from __future__ import annotations

import math
import re
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "morphz_structured_cognitive_transition_training_patent_application_v1.md"
ASSET_DIR = ROOT / "assets"
OUTPUT = ROOT / "Morphz_结构化认知状态转换训练_发明专利申请文件_v1.docx"
INVENTION_TITLE = "一种结构化认知状态转换数据构造及模型训练方法及系统"

# Use the same PostScript font name as the visually approved first patent v6.
# The macOS display name "宋体-简" is not resolved reliably by LibreOffice and
# can render Chinese text as empty boxes during PDF/XML-oriented conversion.
CJK_FONT = "STSong"
LATIN_FONT = "STSong"
MONO_FONT = "STSong"
BLACK = RGBColor(0, 0, 0)
GRAY = RGBColor(100, 100, 100)


def set_run_font(run, size: float, *, bold: bool | None = None,
                 italic: bool | None = None, cjk: str = CJK_FONT,
                 latin: str = LATIN_FONT, color: RGBColor = BLACK) -> None:
    run.font.name = latin
    rfonts = run._element.get_or_add_rPr().rFonts
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)
    rfonts.set(qn("w:eastAsia"), cjk)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style_font(style, size: float, *, bold: bool = False,
                   cjk: str = CJK_FONT, latin: str = LATIN_FONT) -> None:
    style.font.name = latin
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = BLACK
    rfonts = style._element.get_or_add_rPr().rFonts
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)
    rfonts.set(qn("w:eastAsia"), cjk)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text_node = OxmlElement("w:t")
    text_node.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text_node, end])
    set_run_font(run, 9, color=GRAY)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.3)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.4)
    section.header_distance = Cm(1.0)
    section.footer_distance = Cm(1.0)

    styles = doc.styles
    normal = styles["Normal"]
    set_style_font(normal, 11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.28
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    title = styles.add_style("Patent Title", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(title, 22, bold=True)
    title.paragraph_format.space_after = Pt(18)
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.keep_with_next = True

    h1 = styles["Heading 1"]
    set_style_font(h1, 18, bold=True)
    h1.paragraph_format.space_before = Pt(8)
    h1.paragraph_format.space_after = Pt(14)
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h1.paragraph_format.keep_with_next = True

    h2 = styles["Heading 2"]
    set_style_font(h2, 14, bold=True)
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(7)
    h2.paragraph_format.keep_with_next = True

    h3 = styles["Heading 3"]
    set_style_font(h3, 12, bold=True)
    h3.paragraph_format.space_before = Pt(9)
    h3.paragraph_format.space_after = Pt(5)
    h3.paragraph_format.keep_with_next = True

    body = styles.add_style("Patent Body", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(body, 11)
    body.paragraph_format.space_after = Pt(5)
    body.paragraph_format.line_spacing = 1.28
    body.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    body.paragraph_format.widow_control = True

    claims = styles.add_style("Claim Body", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(claims, 11)
    claims.paragraph_format.space_after = Pt(6)
    claims.paragraph_format.line_spacing = 1.32
    claims.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    claims.paragraph_format.widow_control = True

    code = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(code, 8.8, cjk=CJK_FONT, latin=MONO_FONT)
    code.paragraph_format.left_indent = Cm(0.6)
    code.paragraph_format.right_indent = Cm(0.3)
    code.paragraph_format.space_before = Pt(3)
    code.paragraph_format.space_after = Pt(3)
    code.paragraph_format.line_spacing = 1.05

    caption = styles.add_style("Figure Caption", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(caption, 10)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(5)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_field(footer)


def diagram_font(size: float = 11, bold: bool = False):
    candidates = [
        Path("/System/Library/Fonts/Supplemental/Songti.ttc"),
        Path("/System/Library/Fonts/PingFang.ttc"),
        Path("/System/Library/Fonts/STHeiti Medium.ttc"),
    ]
    chosen = next((path for path in candidates if path.exists()), None)
    if chosen is None:
        return ImageFont.load_default()
    index = 1 if bold and chosen.name == "Songti.ttc" else 0
    return ImageFont.truetype(str(chosen), int(size * 2.15), index=index)


class Diagram:
    def __init__(self, width: int = 2000, height: int = 1400):
        self.width = width
        self.height = height
        self.image = Image.new("RGB", (width, height), "white")
        self.draw = ImageDraw.Draw(self.image)
        self.sx = width / 10.0
        self.sy = height / 7.0

    def point(self, x: float, y: float) -> tuple[float, float]:
        return x * self.sx, self.height - y * self.sy

    def rectangle(self, x: float, y: float, w: float, h: float, width: int = 3) -> None:
        left, bottom = self.point(x, y)
        right, top = self.point(x + w, y + h)
        self.draw.rectangle((left, top, right, bottom), outline="black", fill="white", width=width)

    def text(self, x: float, y: float, text: str, *, size: float = 11,
             bold: bool = False) -> None:
        px, py = self.point(x, y)
        self.draw.multiline_text(
            (px, py), text, fill="black", font=diagram_font(size, bold),
            anchor="mm", align="center", spacing=7,
        )

    def line(self, start, end, width: int = 3) -> None:
        x1, y1 = self.point(*start)
        x2, y2 = self.point(*end)
        self.draw.line((x1, y1, x2, y2), fill="black", width=width)

    def arrow(self, start, end, width: int = 3) -> None:
        x1, y1 = self.point(*start)
        x2, y2 = self.point(*end)
        self.draw.line((x1, y1, x2, y2), fill="black", width=width)
        angle = math.atan2(y2 - y1, x2 - x1)
        head = 16
        for offset in (2.55, -2.55):
            hx = x2 + head * math.cos(angle + offset)
            hy = y2 + head * math.sin(angle + offset)
            self.draw.line((x2, y2, hx, hy), fill="black", width=width)

    def polyline_arrow(self, points, width: int = 3) -> None:
        pixels = [self.point(*point) for point in points]
        self.draw.line(pixels, fill="black", width=width, joint="curve")
        (x1, y1), (x2, y2) = pixels[-2:]
        angle = math.atan2(y2 - y1, x2 - x1)
        head = 16
        for offset in (2.55, -2.55):
            hx = x2 + head * math.cos(angle + offset)
            hy = y2 + head * math.sin(angle + offset)
            self.draw.line((x2, y2, hx, hy), fill="black", width=width)

    def save(self, path: Path) -> None:
        self.image.save(path, dpi=(220, 220), optimize=True)


def box(ax: Diagram, x: float, y: float, w: float, h: float,
        label: str, number: str | None = None, size: float = 10.5) -> None:
    ax.rectangle(x, y, w, h)
    text = f"{label}\n{number}" if number else label
    ax.text(x + w / 2, y + h / 2, text, size=size)


def fig1(path: Path) -> None:
    ax = Diagram()
    box(ax, 0.15, 5.35, 1.65, 0.85, "权威事件存储", "101", 9)
    box(ax, 0.15, 3.85, 1.65, 0.85, "认知状态存储", "102", 9)
    box(ax, 0.15, 2.35, 1.65, 0.85, "确定性运行时", "103", 9)
    box(ax, 0.15, 0.85, 1.65, 0.85, "外部资源", "104", 9)
    boxes = [
        (2.25, 5.35, "权威事实获取单元", "110"),
        (4.15, 5.35, "轨迹定界单元", "120"),
        (6.05, 5.35, "因果图构造单元", "130"),
        (7.95, 5.35, "状态视图重建单元", "140"),
        (2.25, 3.35, "转换关联单元", "150"),
        (4.15, 3.35, "训练片段派生单元", "160"),
        (6.05, 3.35, "训练掩码生成单元", "170"),
        (7.95, 3.35, "校验与权限控制单元", "180"),
    ]
    for x, y, label, number in boxes:
        box(ax, x, y, 1.65, 0.85, label, number, 8.5)
    box(ax, 4.15, 1.15, 2.0, 0.9, "模型训练单元", "190", 9.5)
    box(ax, 7.25, 1.15, 1.8, 0.9, "模型存储", "105", 9.5)
    ax.arrow((1.8, 5.75), (2.25, 5.75))
    ax.polyline_arrow([(1.8, 4.25), (2.02, 4.25), (2.02, 5.55), (2.25, 5.55)], 2)
    ax.polyline_arrow([(1.8, 2.75), (2.02, 2.75), (2.02, 5.35), (2.25, 5.35)], 2)
    ax.polyline_arrow([(1.8, 1.25), (2.02, 1.25), (2.02, 5.15), (2.25, 5.15)], 2)
    for x1, x2 in ((3.9, 4.15), (5.8, 6.05), (7.7, 7.95)):
        ax.arrow((x1, 5.75), (x2, 5.75))
    ax.polyline_arrow([(8.78, 5.35), (8.78, 4.7), (3.08, 4.7), (3.08, 4.2)])
    for x1, x2 in ((3.9, 4.15), (5.8, 6.05), (7.7, 7.95)):
        ax.arrow((x1, 3.75), (x2, 3.75))
    ax.polyline_arrow([(8.78, 3.35), (8.78, 2.55), (5.15, 2.55), (5.15, 2.05)])
    ax.arrow((6.15, 1.6), (7.25, 1.6))
    ax.save(path)


def fig2(path: Path) -> None:
    ax = Diagram()
    box(ax, 3.2, 5.8, 3.6, 0.65, "结构化认知状态转换", "200", 12)
    first = [
        (0.15, "基础状态视图", "210"), (2.15, "读取集合", "220"),
        (4.15, "目标及策略绑定", "230"), (6.15, "候选转换", "240"),
        (8.15, "运行时接纳", "250"),
    ]
    second = [
        (0.15, "现实效果", "260"), (2.15, "状态差异及\n后继状态", "270"),
        (4.15, "现实结果", "280"), (6.15, "验证结果", "290"),
        (8.15, "奖励记录", "295"),
    ]
    for x, label, number in first:
        box(ax, x, 3.8, 1.7, 0.85, label, number, 9)
        ax.arrow((5.0, 5.8), (x + 0.85, 4.65), 2)
    for x, label, number in second:
        box(ax, x, 1.25, 1.7, 0.85, label, number, 9)
    for i in range(4):
        ax.arrow((first[i][0] + 1.7, 4.225), (first[i + 1][0], 4.225), 2)
    ax.polyline_arrow([(9.0, 3.8), (9.0, 3.0), (1.0, 3.0), (1.0, 2.1)], 2)
    for i in range(4):
        ax.arrow((second[i][0] + 1.7, 1.675), (second[i + 1][0], 1.675), 2)
    ax.save(path)


def fig3(path: Path) -> None:
    ax = Diagram(width=1500, height=2000)
    steps = [
        ("获取有界执行事实集合", "S301"),
        ("构造类型化因果轨迹", "S302"),
        ("重建决策时状态视图", "S303"),
        ("关联提议、接纳、效果和提交", "S304"),
        ("形成训练输入、目标和后决策信息", "S305"),
        ("生成训练掩码", "S306"),
        ("执行因果、版本、权限和完整性校验", "S307"),
        ("输出训练片段", "S308"),
    ]
    ys = [6.05 - index * 0.76 for index in range(len(steps))]
    for (label, number), y in zip(steps, ys):
        box(ax, 1.7, y, 6.6, 0.52, label, number, 9.6)
    for y1, y2 in zip(ys, ys[1:]):
        ax.arrow((5.0, y1), (5.0, y2 + 0.52), 2)
    ax.save(path)


def fig4(path: Path) -> None:
    ax = Diagram()
    box(ax, 3.8, 5.8, 2.4, 0.7, "基础状态版本", "V42", 10)
    box(ax, 0.5, 4.25, 1.7, 0.75, "线程A决策", "D-A", 9)
    box(ax, 7.8, 4.25, 1.7, 0.75, "线程B决策", "D-B", 9)
    box(ax, 0.3, 2.65, 1.55, 0.75, "执行尝试", "A1失败", 9)
    box(ax, 2.2, 2.65, 1.55, 0.75, "重试执行", "A2成功", 9)
    box(ax, 7.9, 2.65, 1.55, 0.75, "状态提交", "C-B", 9)
    box(ax, 3.9, 0.95, 2.2, 0.8, "汇合节点", "J", 10)
    ax.arrow((4.4, 5.8), (1.35, 5.0))
    ax.arrow((5.6, 5.8), (8.65, 5.0))
    ax.arrow((1.35, 4.25), (1.075, 3.4))
    ax.arrow((1.85, 3.025), (2.2, 3.025))
    ax.arrow((8.65, 4.25), (8.675, 3.4))
    ax.arrow((2.975, 2.65), (4.6, 1.75))
    ax.arrow((8.675, 2.65), (5.4, 1.75))
    ax.text(1.95, 3.48, "重试", size=8.5)
    ax.text(5.0, 2.2, "显式因果依赖", size=8.5)
    ax.save(path)


def fig5(path: Path) -> None:
    ax = Diagram()
    chain = [
        (0.25, "语言模型候选提议", "P"),
        (2.25, "确定性运行时接纳", "A"),
        (4.25, "执行尝试", "E"),
        (6.25, "现实效果回执", "R"),
        (8.25, "权威状态提交", "C"),
    ]
    for x, label, number in chain:
        box(ax, x, 4.1, 1.5, 0.9, label, number, 8.8)
    for current, following in zip(chain, chain[1:]):
        ax.arrow((current[0] + 1.5, 4.55), (following[0], 4.55))
    box(ax, 2.25, 1.55, 1.5, 0.9, "拒绝记录", "X", 9)
    box(ax, 6.25, 1.55, 1.5, 0.9, "提交冲突", "CF", 9)
    ax.arrow((3.0, 4.1), (3.0, 2.45))
    ax.arrow((9.0, 4.1), (7.0, 2.45))
    ax.text(3.7, 3.2, "校验未通过", size=8.5)
    ax.text(7.95, 3.2, "版本失效", size=8.5)
    ax.save(path)


def fig6(path: Path) -> None:
    ax = Diagram()
    box(ax, 0.5, 4.7, 2.0, 0.9, "现实效果回执", "R-401", 9.5)
    box(ax, 4.0, 4.7, 2.0, 0.9, "现实结果", "OUT-501", 9.5)
    box(ax, 7.5, 4.7, 2.0, 0.9, "验证结果", "VER-601", 9.5)
    box(ax, 7.5, 1.75, 2.0, 0.9, "奖励记录", "RW-701", 9.5)
    box(ax, 4.0, 1.75, 2.0, 0.9, "验证器及版本", "V-3", 9.5)
    box(ax, 0.5, 1.75, 2.0, 0.9, "奖励策略及版本", "RP-2", 9.5)
    ax.arrow((2.5, 5.15), (4.0, 5.15))
    ax.arrow((6.0, 5.15), (7.5, 5.15))
    ax.arrow((5.0, 2.65), (8.0, 4.7))
    ax.arrow((8.5, 4.7), (8.5, 2.65))
    ax.polyline_arrow([(2.5, 2.2), (2.8, 1.15), (8.5, 1.15), (8.5, 1.75)], 2)
    ax.text(5.0, 3.95, "事实与判定分离", size=9)
    ax.text(5.6, 0.92, "按奖励策略派生", size=9)
    ax.save(path)


def fig7(path: Path) -> None:
    ax = Diagram()
    box(ax, 3.2, 5.8, 3.6, 0.65, "训练片段", "700", 12)
    regions = [
        (0.15, "输入区", "710"), (1.8, "目标区", "720"),
        (3.45, "环境输出区", "730"), (5.1, "评价区", "740"),
        (6.75, "元数据区", "750"), (8.4, "排除区", "760"),
    ]
    for x, label, number in regions:
        box(ax, x, 3.55, 1.45, 0.85, label, number, 8.8)
        ax.arrow((5.0, 5.8), (x + 0.725, 4.4), 2)
    box(ax, 2.1, 1.15, 5.8, 0.9,
        "训练掩码：模型输入／监督目标／环境输出／评价／元数据／排除", "770", 9.2)
    for x, _, _ in regions:
        ax.arrow((x + 0.725, 3.55), (5.0, 2.05), 2)
    ax.save(path)


def fig8(path: Path) -> None:
    ax = Diagram()
    boxes = [
        (0.15, 4.7, "训练数据存储", "810"),
        (2.05, 4.7, "样本选择单元", "820"),
        (3.95, 4.7, "序列化与掩码单元", "830"),
        (5.85, 4.7, "模型训练单元", "840"),
        (7.75, 4.7, "更新后的模型", "850"),
        (7.75, 2.4, "评测单元", "860"),
        (5.85, 2.4, "部署单元", "870"),
        (3.95, 2.4, "智能体运行时", "880"),
    ]
    for x, y, label, number in boxes:
        box(ax, x, y, 1.6, 0.85, label, number, 8.5)
    for i in range(4):
        ax.arrow((boxes[i][0] + 1.6, 5.125), (boxes[i + 1][0], 5.125))
    ax.arrow((8.55, 4.7), (8.55, 3.25))
    ax.arrow((7.75, 2.825), (7.45, 2.825))
    ax.arrow((5.85, 2.825), (5.55, 2.825))
    ax.polyline_arrow([(3.95, 2.825), (1.0, 2.825), (1.0, 4.7)])
    ax.text(2.5, 2.55, "产生新的权威运行事实", size=8.8)
    ax.save(path)


def build_figures() -> None:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    builders = [fig1, fig2, fig3, fig4, fig5, fig6, fig7, fig8]
    names = [
        "training_patent_fig1_system.png",
        "training_patent_fig2_transition.png",
        "training_patent_fig3_method.png",
        "training_patent_fig4_causal.png",
        "training_patent_fig5_authority.png",
        "training_patent_fig6_verification.png",
        "training_patent_fig7_mask.png",
        "training_patent_fig8_training.png",
    ]
    for builder, name in zip(builders, names):
        builder(ASSET_DIR / name)


def add_plain_paragraph(doc: Document, text: str, claims_mode: bool) -> None:
    style = "Claim Body" if claims_mode else "Patent Body"
    paragraph = doc.add_paragraph(style=style)
    paragraph.paragraph_format.keep_together = False
    paragraph.paragraph_format.keep_with_next = False
    run = paragraph.add_run(text)
    set_run_font(run, 11)


def add_code_block(doc: Document, lines: list[str]) -> None:
    paragraph = doc.add_paragraph(style="Code Block")
    for index, line in enumerate(lines):
        if index:
            paragraph.add_run().add_break()
        run = paragraph.add_run(line)
        set_run_font(run, 8.8, cjk=CJK_FONT, latin=MONO_FONT)
    paragraph.paragraph_format.keep_together = True


def add_image(doc: Document, rel_path: str, alt: str, *, page_break_before: bool = False) -> None:
    path = ROOT / rel_path
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.page_break_before = page_break_before
    width = Inches(5.35 if "图3 " in alt else 6.15)
    picture = paragraph.add_run().add_picture(str(path), width=width)
    picture._inline.docPr.set("title", alt)
    picture._inline.docPr.set("descr", alt)
    caption = doc.add_paragraph(style="Figure Caption")
    caption.add_run(alt)


def build_docx() -> None:
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    doc = Document()
    configure_document(doc)
    in_code = False
    code_lines: list[str] = []
    claims_mode = False
    figures_mode = False
    image_count = 0
    pending_page_break = False

    for line in lines:
        stripped = line.strip()
        if stripped.startswith("```"):
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            continue
        if in_code:
            code_lines.append(line)
            continue
        if not stripped:
            continue
        if stripped == "---PAGEBREAK---":
            pending_page_break = True
            continue

        image_match = re.fullmatch(r"!\[(.+)]\((.+)\)", stripped)
        if image_match:
            alt, rel_path = image_match.groups()
            add_image(doc, rel_path, alt, page_break_before=image_count > 0)
            image_count += 1
            continue

        if stripped.startswith("# "):
            heading = stripped[2:].strip()
            if heading == "说明书":
                claims_mode = False
                figures_mode = False
                continue
            paragraph = doc.add_paragraph(style="Heading 1")
            paragraph.paragraph_format.page_break_before = pending_page_break
            pending_page_break = False
            run = paragraph.add_run(heading)
            set_run_font(run, 18, bold=True)
            claims_mode = heading == "权利要求书"
            figures_mode = heading == "说明书附图"
            if figures_mode:
                image_count = 0
            continue

        if stripped.startswith("## "):
            heading = stripped[3:].strip()
            if heading == INVENTION_TITLE:
                paragraph = doc.add_paragraph(style="Patent Title")
                run = paragraph.add_run(heading)
                set_run_font(run, 22, bold=True)
                continue
            paragraph = doc.add_paragraph(style="Heading 2")
            run = paragraph.add_run(heading)
            set_run_font(run, 14, bold=True)
            continue

        if stripped.startswith("### "):
            paragraph = doc.add_paragraph(style="Heading 3")
            run = paragraph.add_run(stripped[4:].strip())
            set_run_font(run, 12, bold=True)
            continue

        add_plain_paragraph(doc, stripped, claims_mode)

    doc.core_properties.title = INVENTION_TITLE
    doc.core_properties.subject = ""
    doc.core_properties.author = ""
    doc.core_properties.keywords = ""
    doc.core_properties.comments = ""
    doc.save(OUTPUT)


def main() -> None:
    if not SOURCE.exists():
        raise SystemExit(f"missing source: {SOURCE}")
    build_figures()
    build_docx()
    print(OUTPUT)


if __name__ == "__main__":
    main()
